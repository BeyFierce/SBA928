from io import BytesIO

from docx import Document
from pydantic import ValidationError

from sales_agent.demo import demo_brief, demo_findings
from sales_agent.documents import extract_product_document
from sales_agent.prompts import RESEARCH_GUARDRAIL
from sales_agent.quality import evaluate_brief
from sales_agent.research import WebResearcher
from sales_agent.rendering import brief_to_markdown
from sales_agent.schemas import EvidenceItem
from sales_agent.schemas import SalesRequest


def sample_request() -> SalesRequest:
    return SalesRequest(
        product_name="SecureFlow Cloud",
        company_url="https://example.com",
        product_category="Cloud security",
        competitor_urls=["https://competitor.example.com"],
        research_urls=["https://example.com/press-release"],
        value_proposition="Reduce investigation time with unified security visibility.",
        target_customer="CISO",
    )


def test_request_validation():
    request = sample_request()
    assert request.company_url.host == "example.com"


def test_invalid_url_rejected():
    try:
        SalesRequest(
            product_name="Test",
            company_url="not-a-url",
            product_category="Security",
            value_proposition="A useful value proposition",
            target_customer="CISO",
        )
    except ValidationError:
        return
    raise AssertionError("Invalid URL should fail validation")


def test_local_url_blocked():
    try:
        WebResearcher._validate_url("http://localhost:8501")
    except ValueError:
        return
    raise AssertionError("Local URLs should be blocked")


def test_private_ip_url_blocked():
    try:
        WebResearcher._validate_url("http://10.0.0.8/report")
    except ValueError:
        return
    raise AssertionError("Private IP URLs should be blocked")


def test_demo_pipeline():
    request = sample_request()
    finding = demo_findings(request, str(request.company_url), "Company Analyst")
    brief = demo_brief(request, [finding])
    assert brief.prospect_name == "example.com"
    assert brief.discovery_questions


def test_txt_product_document_extraction():
    text = extract_product_document(
        b"SecureFlow Cloud\nAutomates response workflows.", "overview.txt"
    )
    assert "Automates response workflows" in text


def test_docx_product_document_extraction():
    document = Document()
    document.add_heading("SecureFlow Cloud", 1)
    document.add_paragraph("Unifies cloud-security evidence.")
    buffer = BytesIO()
    document.save(buffer)
    text = extract_product_document(buffer.getvalue(), "overview.docx")
    assert "Unifies cloud-security evidence" in text


def test_unsupported_product_document_rejected():
    try:
        extract_product_document(b"binary", "overview.pptx")
    except ValueError:
        return
    raise AssertionError("Unsupported uploads should fail validation")


def test_one_page_renderer_and_quality_review():
    request = sample_request()
    finding = demo_findings(request, str(request.company_url), "Company Analyst")
    brief = demo_brief(request, [finding])
    report = brief_to_markdown(request, brief)
    quality = evaluate_brief(brief)
    assert "# Account Brief:" in report
    assert "## Sources" in report
    assert quality.completeness_percent == 100
    assert quality.status == "Ready for human review"
    assert quality.low_confidence_sources == 1


def test_confidence_label_is_constrained():
    try:
        EvidenceItem(
            claim="Example claim",
            source_url="https://example.com",
            source_title="Example",
            confidence="uncertain",
        )
    except ValidationError:
        return
    raise AssertionError("Confidence must be high, medium, or low")


def test_research_prompt_keeps_evidence_guardrails():
    normalized_prompt = " ".join(RESEARCH_GUARDRAIL.split())
    required_phrases = [
        "Use only the supplied source text",
        "untrusted data",
        "Every factual claim must cite",
        "facts and hypotheses separate",
    ]
    for phrase in required_phrases:
        assert phrase in normalized_prompt
