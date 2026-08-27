from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".docx", ".md", ".pdf", ".txt"}


def extract_product_document(
    file_bytes: bytes,
    filename: str,
    *,
    max_chars: int = 20_000,
) -> str:
    """Extract bounded text from a product overview uploaded by the user."""

    if not file_bytes:
        raise ValueError("The uploaded product document is empty.")

    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported document type. Use one of: {supported}.")

    if extension in {".txt", ".md"}:
        text = file_bytes.decode("utf-8", errors="replace")
    elif extension == ".pdf":
        reader = PdfReader(BytesIO(file_bytes))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    else:
        document = Document(BytesIO(file_bytes))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        text = "\n".join([*paragraphs, *table_cells])

    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not normalized:
        raise ValueError("No readable text was found in the uploaded document.")
    return normalized[:max_chars]
